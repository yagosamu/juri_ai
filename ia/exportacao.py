import re
from io import BytesIO
from django.http import HttpResponse


# ── Helpers de markdown ──────────────────────────────────────────────────────────

def _strip_markup(text: str) -> str:
    """Remove ** e __ do texto (para headings no docx)."""
    return re.sub(r'\*\*|__', '', text).strip()


def _add_runs(paragraph, text: str):
    """Divide a linha em runs, aplicando negrito nos trechos **...**."""
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _md_to_rl(text: str) -> str:
    """Converte **bold** e __bold__ para XML do reportlab (<b>...</b>)."""
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__',     r'<b>\1</b>', text)
    return text


# ── DOCX ─────────────────────────────────────────────────────────────────────────

def gerar_docx(conteudo_markdown: str, nome_arquivo: str) -> HttpResponse:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(90)
        section.right_margin  = Pt(90)

    for line in conteudo_markdown.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph()
        elif s.startswith('### '):
            doc.add_heading(_strip_markup(s[4:]), level=3)
        elif s.startswith('## '):
            doc.add_heading(_strip_markup(s[3:]), level=2)
        elif s.startswith('# '):
            doc.add_heading(_strip_markup(s[2:]), level=1)
        elif s in ('---', '___', '***'):
            p = doc.add_paragraph('─' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif s.startswith('- ') or s.startswith('* '):
            _add_runs(doc.add_paragraph(style='List Bullet'), s[2:])
        elif re.match(r'^\d+\. ', s):
            _add_runs(doc.add_paragraph(style='List Number'), re.sub(r'^\d+\. ', '', s))
        else:
            _add_runs(doc.add_paragraph(), s)

    buf = BytesIO()
    doc.save(buf)
    resp = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    resp['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.docx"'
    return resp


# ── PDF ──────────────────────────────────────────────────────────────────────────

def gerar_pdf(conteudo_markdown: str, nome_arquivo: str) -> HttpResponse:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=2.5*cm, leftMargin=2.5*cm,
                            topMargin=2.5*cm,   bottomMargin=2.5*cm)

    base = getSampleStyleSheet()

    s_title = ParagraphStyle('JTitle', parent=base['Title'],
                             fontSize=14, spaceAfter=10,
                             textColor=colors.HexColor('#1e293b'))
    s_h2    = ParagraphStyle('JH2', parent=base['Heading2'],
                             fontSize=12, spaceBefore=12, spaceAfter=6,
                             textColor=colors.HexColor('#334155'))
    s_h3    = ParagraphStyle('JH3', parent=base['Heading3'],
                             fontSize=11, spaceBefore=8, spaceAfter=4,
                             textColor=colors.HexColor('#475569'))
    s_body  = ParagraphStyle('JBody', parent=base['Normal'],
                             fontSize=11, leading=16, spaceAfter=6,
                             alignment=TA_JUSTIFY)
    s_list  = ParagraphStyle('JList', parent=base['Normal'],
                             fontSize=11, leading=16, leftIndent=20, spaceAfter=4)

    story = []
    for line in conteudo_markdown.splitlines():
        s = line.strip()
        if not s:
            story.append(Spacer(1, 0.3*cm))
        elif s in ('---', '___', '***'):
            story.append(HRFlowable(width='100%', thickness=0.5,
                                    color=colors.HexColor('#cbd5e1'),
                                    spaceBefore=4, spaceAfter=4))
        elif s.startswith('### '):
            story.append(Paragraph(_strip_markup(s[4:]), s_h3))
        elif s.startswith('## '):
            story.append(Paragraph(_strip_markup(s[3:]), s_h2))
        elif s.startswith('# '):
            story.append(Paragraph(_strip_markup(s[2:]), s_title))
        elif s.startswith('- ') or s.startswith('* '):
            story.append(Paragraph(f'• {_md_to_rl(s[2:])}', s_list))
        elif re.match(r'^\d+\. ', s):
            num  = re.match(r'^(\d+)\.', s).group(1)
            text = _md_to_rl(re.sub(r'^\d+\. ', '', s))
            story.append(Paragraph(f'{num}. {text}', s_list))
        else:
            story.append(Paragraph(_md_to_rl(s), s_body))

    doc.build(story)
    resp = HttpResponse(buf.getvalue(), content_type='application/pdf')
    resp['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.pdf"'
    return resp
